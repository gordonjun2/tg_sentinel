import ssl
import urllib3
import requests
import difflib
import numpy as np
from huggingface_hub import configure_http_backend
import markdown
from weasyprint import HTML
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ssl._create_default_https_context = ssl._create_unverified_context

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


# [OPTIONAL] Use if there is SSL certificate verification issues
def backend_factory() -> requests.Session:
    session = requests.Session()
    session.verify = False
    return session


configure_http_backend(backend_factory=backend_factory)

sys_msg = """
I will provide you with a segment of a transcription from a spoken discussion. The transcription may not be perfect, but please do your best to make sense of it.

Please analyze the content and extract the following:

------
Key Discussion Points:
List the main points or topics being discussed in this segment. Capture important ideas, arguments, or themes. Keep them concise but informative.

Mentioned Projects / Products / Tools:
Identify all specific projects, products, tools, companies, or technologies mentioned in the text. For each one, provide a short description based on the context, or from external knowledge if the context is unclear or transcription appears inaccurate.
Some names or terms may be misspelled. Use your knowledge of known technologies, companies, and tools to infer the most likely intended references and replace or correct them as needed. Aim for accuracy and relevance.

Notable Quotes or Opinions:
Extract any statements that reflect strong opinions, insightful comments, or notable quotes. Attribute them to the speaker if speaker names or labels are present.
------

The transcript will be sent in batches, so treat each batch as part of a larger conversation. Do not summarize or conclude; just extract structured information for the current segment.
"""

sys_msg_final_summary = """
You will be provided with a series of structured insights extracted from multiple transcription segments. Your task is to combine them into a single, cohesive markdown document that summarizes the discussion as a whole.

Please follow this structure exactly:

------
# Executive Summary
Provide a concise overview of the key themes and takeaways from the entire discussion. Focus on overarching insights, decisions, and purpose of the conversation.

# Key Discussion Points
Combine and deduplicate the main discussion points from all segments. Group similar ideas and arrange them logically to reflect the flow of the conversation. Use bullet points and preserve clarity.

# Mentioned Projects / Products / Tools
List all unique projects, products, tools, companies, or technologies that were mentioned across segments. If any were repeated, merge context where relevant. Include short descriptions if they were provided.

# Notable Quotes or Opinions (optional)
Include a curated list of the most insightful or opinionated quotes. Attribute them if speaker names or labels were available. Only include quotes that add meaningful perspective.
------

Guidelines:
- Do not summarize the original paragraphs again. Only work with the structured outputs you were given.
- Ensure **no key points or tools mentioned are lost** during merging.
- Focus on **clarity, conciseness, and logical organization**.
- Output only markdown content — do not explain or add commentary.
- Do not add the '------' separator at the start and the end of the output.

The final result should be ready for presentation in a document or report.
"""


def chunk_audio_with_overlap(audio_array,
                             chunk_size_seconds=30,
                             overlap_seconds=3,
                             sample_rate=16000):
    """Split audio into overlapping chunks."""
    chunk_size = chunk_size_seconds * sample_rate
    overlap_size = overlap_seconds * sample_rate
    step = chunk_size - overlap_size

    chunks = []
    for i in range(0, len(audio_array), step):
        chunk = audio_array[i:i + chunk_size]
        if len(chunk) < chunk_size:
            chunk = np.pad(chunk, (0, chunk_size - len(chunk)))
        chunks.append(chunk)
    return chunks


def remove_overlap_text(prev_text,
                        curr_text,
                        window_words=20,
                        similarity_threshold=0.85):
    """
    Remove overlapping part from the start of curr_text that duplicates the end of prev_text,
    based on longest matching sequence of words.
    
    Args:
        prev_text (str): Previous chunk's full transcription.
        curr_text (str): Current chunk's full transcription.
        window_words (int): Number of words to consider at the overlap boundary.
        similarity_threshold (float): Minimum similarity to consider a match.
        
    Returns:
        str: Current chunk's transcription with overlapping start removed.
    """

    prev_words = prev_text.split()
    curr_words = curr_text.split()

    # Get last window_words from prev_text and first window_words from curr_text
    prev_tail = prev_words[-window_words:]
    curr_head = curr_words[:window_words]

    # Find longest matching sequence of words at the boundary
    # We'll try decreasing length sequences from window_words down to 1,
    # stopping at the longest matching sequence with similarity >= threshold.

    for overlap_len in range(window_words, 0, -1):
        prev_sub = prev_tail[-overlap_len:]
        curr_sub = curr_head[:overlap_len]

        # Compute similarity between word sequences (join to string and use difflib)
        prev_sub_str = " ".join(prev_sub)
        curr_sub_str = " ".join(curr_sub)
        ratio = difflib.SequenceMatcher(None, prev_sub_str,
                                        curr_sub_str).ratio()

        if ratio >= similarity_threshold:
            # Remove the overlapping words from the start of curr_text
            # Rejoin words skipping the overlap
            return " ".join(curr_words[overlap_len:]).lstrip()

    # No sufficient overlap found, return original curr_text
    return curr_text


def convert_text_to_docx(markdown_text, output_file=None):
    if output_file is None:
        # Get current timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Create filename without double underscore
        output_file = f"./discussion_insights/discussion_insights_{timestamp}.docx"

    # Create a new Document
    doc = Document()

    # Split the markdown text into lines
    lines = markdown_text.split('\n')

    for line in lines:
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        # Handle bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
        # Handle regular text
        elif line.strip():
            p = doc.add_paragraph(line.strip())

    # Save the document
    doc.save(output_file)


def convert_text_to_md(markdown_text, output_file=None):
    if output_file is None:
        # Get current timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Create filename without double underscore
        output_file = f"./discussion_insights/discussion_insights_{timestamp}.md"
    with open(output_file, "w", encoding="utf-8") as file:
        file.write(markdown_text)


def convert_text_to_pdf(markdown_text, output_file=None):
    if output_file is None:
        # Get current timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Create filename without double underscore
        output_file = f"./discussion_insights/discussion_insights_{timestamp}.pdf"

    html_text = markdown.markdown(markdown_text, output_format='html5')

    html_document = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{ font-family: Arial, sans-serif; margin: 2em; }}
            h1, h2, h3 {{ color: #333; }}
            ul {{ margin-left: 1em; }}
        </style>
    </head>
    <body>
    {html_text}
    </body>
    </html>
    """

    HTML(string=html_document).write_pdf(output_file)
